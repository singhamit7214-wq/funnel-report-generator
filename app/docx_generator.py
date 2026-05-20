"""Generate a Word document matching the weekly funnel report format."""

import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.models import FunnelReport


def _set_cell(cell, text: str, bold: bool = False, size: int = 8,
              align: str = "center", bg_color: str = None):
    """Helper to format a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bg_color:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = cell._element.makeelement(qn("w:shd"), {
            qn("w:fill"): bg_color,
            qn("w:val"): "clear",
        })
        tc_pr.append(shd)


def _pct_str(val) -> str:
    """Format a percentage value."""
    if val is None:
        return "N/A"
    return f"{val:.1f}%"


def _add_highlight_section(doc: Document, report: FunnelReport):
    """Add the Highlight & Demand Updates section."""
    doc.add_heading("Highlight & Demand Updates:", level=2)

    metrics = report
    summary = (
        f"In {report.report_year}, we've delivered {metrics.ytd_os_inclines} OS inclines "
        f"across L4-L6 roles, leading to {metrics.ytd_offer_accepts} Offer Accepts, "
        f"{metrics.ytd_offer_stage} at Offer Stage and "
        f"{metrics.ytd_offer_declines} Offer Declines."
    )
    doc.add_paragraph(summary)

    # Overall demand paragraph
    demand = (
        f"Overall Demand Update - We are working on {metrics.total_requisitions} "
        f"requisitions with a pipeline of {metrics.total_active_pipeline} candidates "
        f"at advanced stages. The funnel demonstrates steady early-stage throughput "
        f"with BPS-to-Onsite progression at {_pct_str(metrics.bps_to_os_pct)} "
        f"and OS-to-incline conversion at {_pct_str(metrics.os_to_incline_pct)}. "
        f"{metrics.ytd_os_inclines + metrics.ytd_offer_stage + metrics.ytd_offer_accepts + metrics.ytd_offer_declines} "
        f"offers have been released to-date, {metrics.ytd_offer_accepts} of which "
        f"stand accepted and {metrics.ytd_offer_stage} have been extended/in-process, "
        f"reflecting {_pct_str(metrics.offer_acceptance_pct)} offer acceptance rate."
    )
    doc.add_paragraph(demand)


def _add_business_updates(doc: Document, report: FunnelReport):
    """Add the Business-wise updates section."""
    doc.add_heading("Business wise updates:", level=2)

    for idx, biz in enumerate(report.business_updates, 1):
        doc.add_heading(f"{idx}. {biz.business_name}", level=3)

        delivered = (
            f"GSS delivered {biz.os_inclines} OS inclines resulting in "
            f"{biz.offer_accepts} offer accepts"
        )
        if biz.offer_stage:
            delivered += f" and {biz.offer_stage} at Offer Stage"
        if biz.offer_declines:
            delivered += f" and {biz.offer_declines} offer Decline(s)"
        delivered += "."
        doc.add_paragraph(delivered, style="List Bullet")

        if biz.incline_conversion_pct is not None:
            conv = (
                f"Our onsite incline conversion currently stands at "
                f"{biz.incline_conversion_pct:.0f}% "
                f"({biz.os_inclines}/{biz.os_completed}) against the goal of "
                f"{biz.conversion_goal_pct:.0f}%"
            )
            pipeline_parts = []
            if biz.pipeline_onsite:
                pipeline_parts.append(f"{biz.pipeline_onsite} at Onsite")
            if biz.pipeline_bps:
                pipeline_parts.append(f"{biz.pipeline_bps} BPS stage")
            if biz.pipeline_assessment:
                pipeline_parts.append(f"{biz.pipeline_assessment} at Assessment")
            if biz.pipeline_hm_review:
                pipeline_parts.append(f"{biz.pipeline_hm_review} at HM Review")
            total_pipeline = (
                biz.pipeline_onsite + biz.pipeline_bps
                + biz.pipeline_assessment + biz.pipeline_hm_review
            )
            if pipeline_parts:
                conv += (
                    f", with an active pipeline of {total_pipeline} candidates "
                    f"across interview stages ({', '.join(pipeline_parts)})."
                )
            else:
                conv += "."
            doc.add_paragraph(conv, style="List Bullet")

        if biz.notes:
            doc.add_paragraph(biz.notes, style="List Bullet")


def _add_monthly_glidepath(doc: Document, report: FunnelReport):
    """Add the YTD Month-on-Month delivery glidepath table."""
    if not report.monthly_delivery:
        return

    doc.add_heading("YTD Month-on-Month delivery glidepath", level=2)

    months = report.monthly_delivery
    num_months = len(months)
    # Table: header column + month columns + YTD column
    cols = 2 + num_months  # label col + months + YTD
    table = doc.add_table(rows=7, cols=cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    _set_cell(table.rows[0].cells[0], "FGBS " + str(report.report_year), bold=True, bg_color="4472C4")
    _set_cell(table.rows[0].cells[1], "Month", bold=True, bg_color="4472C4")
    for i, m in enumerate(months):
        _set_cell(table.rows[0].cells[2 + i], m.month, bold=True, bg_color="4472C4")

    # Row labels
    labels = [
        ("", "No Of Sourcer"),
        ("OS Inclines", "OS Inclines Actuals"),
        ("", "Onsite Completed"),
        ("Offer Accept", "Offer Accepts Actuals"),
        ("", "Offer Accept Target"),
        ("", "GAP"),
    ]

    for row_idx, (cat, label) in enumerate(labels):
        _set_cell(table.rows[row_idx + 1].cells[0], cat, bold=bool(cat), align="left")
        _set_cell(table.rows[row_idx + 1].cells[1], label, align="left")

    # Fill data
    ytd_sourcers = 0
    ytd_os = 0
    ytd_onsite = 0
    ytd_oa_actual = 0
    ytd_oa_target = 0

    for i, m in enumerate(months):
        col = 2 + i
        _set_cell(table.rows[1].cells[col], str(m.num_sourcers))
        _set_cell(table.rows[2].cells[col], str(m.os_inclines_actual))
        _set_cell(table.rows[3].cells[col], str(m.onsite_completed))
        _set_cell(table.rows[4].cells[col], str(m.offer_accepts_actual))
        _set_cell(table.rows[5].cells[col], str(m.offer_accept_target))
        gap = m.offer_accept_target - m.offer_accepts_actual
        _set_cell(table.rows[6].cells[col], str(gap))

        ytd_sourcers += m.num_sourcers
        ytd_os += m.os_inclines_actual
        ytd_onsite += m.onsite_completed
        ytd_oa_actual += m.offer_accepts_actual
        ytd_oa_target += m.offer_accept_target

    doc.add_paragraph("")


def _add_pipeline_slate(doc: Document, report: FunnelReport):
    """Add the full pipeline slate table."""
    if not report.skill_groups:
        return

    doc.add_heading("Active Pipeline - YTD Slate", level=2)

    from app.analysis_engine import compute_funnel_metrics
    metrics = compute_funnel_metrics(report.skill_groups)
    t = metrics["totals"]

    # Column headers for the slate table
    headers = [
        "Skill Group", "Demands",
        "RR\nAwait", "RR\nIncl", "RR\nRej",
        "OA\nSched", "OA\nIncl", "OA\nRej",
        "BPS\nAwait", "BPS\nSched", "BPS\nIncl", "BPS\nRej",
        "OS\nAwait", "OS\nSched", "OS\nIncl", "OS\nRej",
        "Offer\nMade", "Offer\nAccept", "Offer\nRenege", "Onboard",
    ]

    num_rows = len(report.skill_groups) + 2  # header + data + grand total
    table = doc.add_table(rows=num_rows, cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for col_idx, h in enumerate(headers):
        _set_cell(table.rows[0].cells[col_idx], h, bold=True, size=7, bg_color="4472C4")

    # Data rows
    fields = [
        "demands",
        "rr_awaiting", "rr_incline", "rr_reject",
        "oa_scheduled", "oa_incline", "oa_reject",
        "bps_awaiting", "bps_scheduled", "bps_incline", "bps_reject",
        "os_awaiting", "os_scheduled", "os_incline", "os_reject",
        "offer_made", "offer_accept", "offer_renege", "onboard",
    ]

    for row_idx, sg in enumerate(report.skill_groups):
        r = row_idx + 1
        _set_cell(table.rows[r].cells[0], sg.skill_group, size=7, align="left")
        for col_idx, field in enumerate(fields):
            val = getattr(sg, field, 0)
            _set_cell(table.rows[r].cells[col_idx + 1], str(val), size=7)

    # Grand total row
    gt_row = num_rows - 1
    _set_cell(table.rows[gt_row].cells[0], "Grand Total", bold=True, size=7, align="left", bg_color="D9E2F3")
    for col_idx, field in enumerate(fields):
        _set_cell(
            table.rows[gt_row].cells[col_idx + 1],
            str(t[field]),
            bold=True, size=7, bg_color="D9E2F3",
        )

    # Conversion metrics row
    doc.add_paragraph("")
    conv_text = (
        f"OA Completed = {metrics['oa_completed']} | "
        f"BPS Completed = {metrics['bps_completed']} | "
        f"OS Completed = {metrics['os_completed']} | "
        f"Offers Released = {metrics['offers_released']}"
    )
    doc.add_paragraph(conv_text)

    conv_detail = (
        f"Conversion: OA Pass {_pct_str(metrics['oa_pass_pct'])} | "
        f"BPS→OS {_pct_str(metrics['bps_to_os_pct'])} | "
        f"OS Incline {_pct_str(metrics['os_to_incline_pct'])} | "
        f"Offer Accept {_pct_str(metrics['offer_accept_pct'])}"
    )
    doc.add_paragraph(conv_detail)


def _add_support_section(doc: Document, report: FunnelReport):
    """Add the Support Required section."""
    if not report.support_items:
        return
    doc.add_heading("Support Required:", level=2)
    for item in report.support_items:
        doc.add_paragraph(item, style="List Bullet")


def _add_ctc_section(doc: Document, report: FunnelReport):
    """Add the Sourcing Insights / CTC section."""
    if not report.ctc_insights:
        return
    doc.add_heading("Sourcing Insights:", level=2)
    doc.add_heading("Capture-the-Conversation - Candidate Analysis", level=3)

    total = report.ctc_insights[0].total if report.ctc_insights else 0
    doc.add_paragraph(
        f"We have captured {total} candidates in {report.report_year} "
        f"with the below reasons documented:"
    )
    for insight in report.ctc_insights:
        pct = insight.percentage
        doc.add_paragraph(
            f"{pct}% ({insight.count}/{insight.total}) owing to {insight.reason}",
            style="List Number",
        )

    # Source mix
    if report.source_hire_pct is not None or report.source_linkedin_pct is not None:
        doc.add_heading("Source Mix:", level=3)
        parts = []
        if report.source_hire_pct is not None:
            parts.append(f"{report.source_hire_pct:.1f}% from Hire")
        if report.source_linkedin_pct is not None:
            parts.append(f"{report.source_linkedin_pct:.1f}% from LinkedIn")
        doc.add_paragraph(
            f"Our YTD pipeline of {report.total_sourced} candidates consists of "
            + " and ".join(parts) + "."
        )


def generate_docx(report: FunnelReport) -> bytes:
    """Generate the complete Word document and return as bytes."""
    doc = Document()

    # Title
    title = doc.add_heading(report.report_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_highlight_section(doc, report)
    _add_business_updates(doc, report)
    _add_monthly_glidepath(doc, report)
    _add_pipeline_slate(doc, report)
    _add_support_section(doc, report)
    _add_ctc_section(doc, report)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
